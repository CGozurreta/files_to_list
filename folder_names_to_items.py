import os
from pathlib import Path
from datetime import datetime
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Pt

def scan_folders(directory, full_file_names, keep_underscores):
    """
    Recursively scan a directory and its subfolders, generating a bulleted list of folders and their contents.
    
    Parameters:
    directory (str): The root directory to scan.
    full_file_names (bool): Whether to use full file names or modified file names.
    keep_underscores (bool): Whether to keep underscores in file names.
    """
    output_file = Path(__file__).parent / "items in folders list.docx"
    document = Document()
    document.add_heading(f"Folder Scan Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=1)
    
    remove_list = ["SFX_AMB_", "SFX_AMB_EP", "EP_", "HH_", "SFX_MG_", "SFX_INT_", "SFX_IT_", "SFX_SHOP_", "TT_", "01.wav", "02.wav", "bloxburg_MS4_UI_", "bburg_desktop_", "bburg_desktop_", "Bloxburg MS4 BG ", "Bloxburg_MS4_"]
    
    try:
        def write_folder(path, level=0):
            """
            Recursively write the folder structure to the output file.
            
            Parameters:
            path (pathlib.Path): The current folder path.
            level (int): The current nesting level (for indentation).
            """
            p = document.add_paragraph(f"{path.name}", style='List Bullet')
            p.paragraph_format.left_indent = Pt(level * 12)  # Indent for nested levels (12 points per level)
            
            for entry in path.glob("*"):
                if entry.is_dir():
                    write_folder(entry, level + 1)
                else:
                    file_name = entry.name
                    if not full_file_names:
                        for remove_str in remove_list:
                            file_name = file_name.replace(remove_str, "")
                    if not keep_underscores:
                        file_name = file_name.replace("_", " ")
                    p = document.add_paragraph(f"{file_name}", style='List Bullet')
                    p.paragraph_format.left_indent = Pt((level + 1) * 12)
        
        write_folder(Path(directory))
        document.save(output_file)
        print(f"Folder scan complete. Output saved to: {output_file}")
    except Exception as e:
        print(f"An error occurred: {e}")

# Example usage
directory = filedialog.askdirectory()
if directory:
    full_file_names = messagebox.askyesno("File Name Option", "Do you want full file names?")
    keep_underscores = messagebox.askyesno("Underscore Option", "Keep _ ?")
    scan_folders(directory, full_file_names, keep_underscores)
else:
    print("No directory selected.")
